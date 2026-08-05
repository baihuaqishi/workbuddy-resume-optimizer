#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
extract_text.py —— 从简历 / JD 文件中提取纯文本。

支持格式：.md / .markdown / .txt / .docx / .pdf
用法：
    python extract_text.py <input_file> [output_file]

若不指定 output_file，文本直接打印到 stdout。
依赖（在隔离 venv 中安装）：python-docx, pypdf
"""
import sys
import os


def extract_md(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_docx(path):
    try:
        from docx import Document
    except ImportError:
        sys.stderr.write("[错误] 未安装 python-docx，请先 pip install python-docx\n")
        sys.exit(2)
    doc = Document(path)
    lines = []
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text)
    # 表格内容
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def extract_pdf(path):
    try:
        from pypdf import PdfReader
    except ImportError:
        sys.stderr.write("[错误] 未安装 pypdf，请先 pip install pypdf\n")
        sys.exit(2)
    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        txt = page.extract_text() or ""
        parts.append(txt)
    return "\n".join(parts)


def extract(path, out=None):
    ext = os.path.splitext(path)[1].lower()
    if ext in (".md", ".markdown"):
        text = extract_md(path)
    elif ext == ".txt":
        text = extract_txt(path)
    elif ext == ".docx":
        text = extract_docx(path)
    elif ext == ".pdf":
        text = extract_pdf(path)
    else:
        sys.stderr.write(f"[错误] 不支持的扩展名：{ext}\n")
        sys.exit(3)

    if out:
        with open(out, "w", encoding="utf-8") as f:
            f.write(text)
        sys.stderr.write(f"[完成] 已提取 {len(text)} 字符 -> {out}\n")
    else:
        sys.stdout.write(text)
    return text


if __name__ == "__main__":
    if len(sys.argv) < 2:
        sys.stderr.write("用法: python extract_text.py <input_file> [output_file]\n")
        sys.exit(1)
    extract(sys.argv[1], sys.argv[2] if len(sys.argv) > 2 else None)
